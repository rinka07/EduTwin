# -*- coding: utf-8 -*-
"""Extraction et découpage des fiches TP (PDF / DOCX).

Deux responsabilités :
1. Extraire le texte brut d'un fichier PDF (via pypdf) ou DOCX (via
   python-docx), détecté par son extension, puis le nettoyer.
2. Découper ce texte en `nb_taches` tâches [{"titre","consigne"}] et en
   chunks pour le RAG.

Le découpage en tâches est robuste : on tente d'abord de repérer des
en-têtes de sections ("Tâche 1", "Exercice 2", "Partie A", "1.") ; à défaut,
on répartit équitablement le texte. On garantit len(taches) == nb_taches.
"""
from __future__ import annotations

import os
import re


class FichierIllisible(Exception):
    """Levée quand un fichier ne peut être lu ou est d'un format non géré."""


# --- Extraction -----------------------------------------------------------
def _extraire_pdf(chemin: str) -> str:
    """Extrait le texte d'un PDF via pypdf, page par page."""
    from pypdf import PdfReader  # import local : lib optionnelle au chargement

    try:
        lecteur = PdfReader(chemin)
    except Exception as exc:  # fichier corrompu, chiffré, etc.
        raise FichierIllisible(f"PDF illisible : {exc}") from exc

    morceaux: list[str] = []
    for page in lecteur.pages:
        try:
            morceaux.append(page.extract_text() or "")
        except Exception:
            # Une page problématique ne doit pas faire échouer tout le fichier.
            continue
    return "\n".join(morceaux)


def _extraire_docx(chemin: str) -> str:
    """Extrait le texte d'un DOCX via python-docx (paragraphes + tableaux)."""
    from docx import Document  # import local : lib optionnelle au chargement

    try:
        document = Document(chemin)
    except Exception as exc:
        raise FichierIllisible(f"DOCX illisible : {exc}") from exc

    morceaux: list[str] = [p.text for p in document.paragraphs]
    # On récupère aussi le texte contenu dans les tableaux.
    for tableau in document.tables:
        for ligne in tableau.rows:
            for cellule in ligne.cells:
                if cellule.text:
                    morceaux.append(cellule.text)
    return "\n".join(morceaux)


def extraire_texte(chemin_fichier: str) -> str:
    """Extrait le texte brut d'un fichier PDF ou DOCX selon son extension."""
    if not os.path.isfile(chemin_fichier):
        raise FichierIllisible(f"Fichier introuvable : {chemin_fichier}")

    extension = os.path.splitext(chemin_fichier)[1].lower()
    if extension == ".pdf":
        texte = _extraire_pdf(chemin_fichier)
    elif extension in (".docx", ".doc"):
        texte = _extraire_docx(chemin_fichier)
    else:
        raise FichierIllisible(
            f"Format non supporté : '{extension}' (PDF ou DOCX attendu)"
        )

    texte = nettoyer_texte(texte)
    if not texte:
        raise FichierIllisible("Le fichier ne contient aucun texte exploitable.")
    return texte


# --- Nettoyage ------------------------------------------------------------
def nettoyer_texte(texte: str) -> str:
    """Normalise les espaces et les sauts de ligne d'un texte extrait."""
    if not texte:
        return ""
    # Uniformise les fins de ligne.
    texte = texte.replace("\r\n", "\n").replace("\r", "\n")
    # Retire les espaces en fin de ligne.
    texte = re.sub(r"[ \t]+\n", "\n", texte)
    # Réduit les espaces multiples à un seul.
    texte = re.sub(r"[ \t]{2,}", " ", texte)
    # Réduit les enchaînements de plus de deux sauts de ligne à deux.
    texte = re.sub(r"\n{3,}", "\n\n", texte)
    return texte.strip()


# --- Découpage en tâches --------------------------------------------------
# Motifs d'en-têtes de sections courants dans les énoncés de TP.
_MOTIF_ENTETE = re.compile(
    r"^\s*(?:"
    r"(?:t[aâ]che|exercice|partie|question|etape|étape|section|q)\s*"
    r"[°ndeos]*\s*\d+[.):\-]?"           # ex. "Tâche 1", "Exercice 2 :"
    r"|\d+\s*[.)]\s+"                       # ex. "1. ", "2) "
    r"|(?:partie|exercice|t[aâ]che)\s+[a-z][.):\-]?"  # ex. "Partie A"
    r")",
    re.IGNORECASE,
)


def _detecter_sections(texte: str) -> list[tuple[str, str]]:
    """Tente de découper le texte en sections via des en-têtes détectés.

    Retourne une liste de tuples (titre, contenu). Vide si rien de fiable.
    """
    lignes = texte.split("\n")
    sections: list[tuple[str, list[str]]] = []
    for ligne in lignes:
        if _MOTIF_ENTETE.match(ligne) and ligne.strip():
            # Nouvel en-tête détecté : on ouvre une section.
            sections.append((ligne.strip(), []))
        elif sections:
            # Ligne de contenu rattachée à la section courante.
            sections[-1][1].append(ligne)
        # (les lignes avant tout en-tête sont ignorées comme préambule)

    resultat: list[tuple[str, str]] = []
    for titre, corps in sections:
        contenu = "\n".join(corps).strip()
        # Le titre lui-même fait partie de la consigne (utile au contexte).
        consigne = (titre + "\n" + contenu).strip() if contenu else titre
        resultat.append((titre, consigne))
    return resultat


def _repartir_equitablement(texte: str, nb_taches: int) -> list[tuple[str, str]]:
    """Répartit le texte en `nb_taches` blocs de taille comparable.

    Découpe par paragraphes puis regroupe pour approcher une taille égale.
    """
    paragraphes = [p.strip() for p in texte.split("\n\n") if p.strip()]
    if not paragraphes:
        paragraphes = [texte.strip()]

    # Si moins de paragraphes que de tâches, on redécoupe par phrases/lignes.
    if len(paragraphes) < nb_taches:
        morceaux = re.split(r"(?<=[.!?])\s+|\n", texte)
        paragraphes = [m.strip() for m in morceaux if m.strip()]
    if not paragraphes:
        paragraphes = [texte.strip() or "(contenu vide)"]

    # Répartition en nb_taches groupes contigus, tailles quasi égales.
    n = len(paragraphes)
    base, reste = divmod(n, nb_taches)
    blocs: list[tuple[str, str]] = []
    debut = 0
    for i in range(nb_taches):
        taille = base + (1 if i < reste else 0)
        # Si base == 0, chaque tâche reçoit au plus un paragraphe (certaines vides).
        part = paragraphes[debut : debut + taille] if taille else []
        debut += taille
        contenu = "\n\n".join(part).strip()
        titre = f"Tâche {i + 1}"
        consigne = contenu if contenu else "(À préciser d'après la fiche TP.)"
        blocs.append((titre, consigne))
    return blocs


def decouper_en_taches(texte: str, nb_taches: int) -> list[dict]:
    """Découpe le texte en exactement `nb_taches` tâches [{"titre","consigne"}].

    Stratégie : détection de sections d'abord, sinon répartition équitable.
    On garantit toujours len(retour) == nb_taches (complétion / troncature).
    """
    if nb_taches < 1:
        nb_taches = 1

    sections = _detecter_sections(texte)

    # On n'utilise la détection que si elle donne un nombre de sections
    # cohérent (au moins la moitié attendue), sinon la répartition est plus sûre.
    if len(sections) >= max(2, nb_taches // 2) and len(sections) >= 1:
        blocs = sections
    else:
        blocs = _repartir_equitablement(texte, nb_taches)

    # Ajustement pour garantir exactement nb_taches éléments.
    if len(blocs) > nb_taches:
        # Trop de sections : on fusionne le surplus dans la dernière tâche.
        tete = blocs[: nb_taches - 1]
        reste = blocs[nb_taches - 1 :]
        titre_fusion = reste[0][0]
        consigne_fusion = "\n\n".join(c for _, c in reste).strip()
        blocs = tete + [(titre_fusion, consigne_fusion)]
    elif len(blocs) < nb_taches:
        # Pas assez : on complète par des tâches génériques.
        for i in range(len(blocs), nb_taches):
            blocs.append(
                (f"Tâche {i + 1}", "(À préciser d'après la fiche TP.)")
            )

    # Normalise en dictionnaires, avec des titres non vides.
    taches: list[dict] = []
    for i, (titre, consigne) in enumerate(blocs, start=1):
        titre_propre = (titre or f"Tâche {i}").strip()
        if len(titre_propre) > 120:  # un titre ne doit pas être un paragraphe
            titre_propre = f"Tâche {i}"
        taches.append(
            {"titre": titre_propre, "consigne": (consigne or "").strip()}
        )
    return taches


# --- Découpage en chunks RAG ---------------------------------------------
def decouper_en_chunks(
    texte: str, taille_cible: int = 600, chevauchement: int = 80
) -> list[str]:
    """Découpe le texte en chunks (~`taille_cible` caractères) pour le RAG.

    On coupe préférentiellement aux frontières de paragraphes, avec un léger
    chevauchement pour ne pas perdre le contexte aux jointures.
    """
    texte = texte.strip()
    if not texte:
        return []

    paragraphes = [p.strip() for p in texte.split("\n\n") if p.strip()]
    if not paragraphes:
        paragraphes = [texte]

    chunks: list[str] = []
    courant = ""
    for para in paragraphes:
        if not courant:
            courant = para
        elif len(courant) + len(para) + 2 <= taille_cible:
            courant += "\n\n" + para
        else:
            chunks.append(courant)
            # Chevauchement : on repart avec la fin du chunk précédent, en
            # démarrant à une frontière de mot pour ne pas couper un mot.
            queue = courant[-chevauchement:] if chevauchement else ""
            if queue:
                espace = queue.find(" ")
                if 0 <= espace < len(queue) - 1:
                    queue = queue[espace + 1 :]
            courant = (queue + "\n\n" + para).strip() if queue else para
    if courant:
        chunks.append(courant)

    # Un très long paragraphe unique peut dépasser la taille cible : on le
    # redécoupe brutalement par fenêtres pour garder des chunks exploitables.
    chunks_finaux: list[str] = []
    for c in chunks:
        if len(c) <= taille_cible * 2:
            chunks_finaux.append(c)
        else:
            pas = taille_cible - chevauchement
            for i in range(0, len(c), pas):
                chunks_finaux.append(c[i : i + taille_cible])
    return [c for c in chunks_finaux if c.strip()]


# --- Point d'entrée du module --------------------------------------------
def extraire_et_decouper(chemin_fichier: str, nb_taches: int) -> dict:
    """Extrait le texte d'un PDF/DOCX puis le découpe en tâches et en chunks.

    Retour : {"taches": [{"titre","consigne"}, ...] (len==nb_taches),
              "chunks": [str, ...]}
    """
    texte = extraire_texte(chemin_fichier)
    taches = decouper_en_taches(texte, nb_taches)
    chunks = decouper_en_chunks(texte)
    return {"taches": taches, "chunks": chunks}
